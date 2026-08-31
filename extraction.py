"""
File-to-text extraction with a confidence gate.

Per the requirements doc: PDF and DOCX only, one file per upload, and
a failed or near-empty extraction must stop the flow and ask for a
re-upload rather than proceeding to feedback. It must never claim to
judge layout/formatting, since only text is available downstream.

Photo stripping is implicit here too: only .extract_text()-style text
is ever pulled out, images are never read, so no photo bytes ever
leave this module.
"""

import io

import pdfplumber
from docx import Document

MIN_CHARS = 120  # confidence gate threshold
MAX_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB cap, matches the requirements doc


class ExtractionResult:
    """
    `meta` carries the document facts that only the FILE knows - page count,
    fonts, whether the layout uses tables, whether images are embedded. The
    format check in the feedback report is built from these rather than guessed
    from the text, because a model reading extracted text genuinely cannot see
    how many pages there are or which font was used.
    """

    def __init__(self, ok: bool, text: str = "", error: str = "", meta: dict | None = None):
        self.ok = ok
        self.text = text
        self.error = error
        self.meta = meta or {}


def _extract_pdf(file_bytes: bytes) -> tuple[str, dict]:
    text_parts, fonts = [], set()
    image_count, table_count = 0, 0
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
            for char in page.chars:
                name = char.get("fontname") or ""
                # PDF font names are usually subset-tagged, e.g. "ABCDEF+Arial-Bold".
                fonts.add(name.split("+")[-1].split("-")[0].split(",")[0])
            image_count += len(page.images or [])
            try:
                table_count += len(page.find_tables())
            except Exception:
                pass  # table detection is best-effort; never fail extraction over it
    meta = {
        "source": "pdf",
        "page_count": page_count,
        "fonts": sorted(f for f in fonts if f),
        "image_count": image_count,
        "table_count": table_count,
    }
    return "\n".join(text_parts), meta


def _extract_docx(file_bytes: bytes) -> tuple[str, dict]:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    fonts = set()
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font is not None and run.font.name:
                fonts.add(run.font.name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    try:
        image_count = len(doc.inline_shapes)
    except Exception:
        image_count = 0
    meta = {
        "source": "docx",
        # Word doesn't store a page count that survives extraction - it's a
        # rendering property. Left as None rather than guessed, so the format
        # check can say "couldn't determine" instead of inventing a number.
        "page_count": None,
        "fonts": sorted(fonts),
        "image_count": image_count,
        "table_count": len(doc.tables),
    }
    return "\n".join(parts), meta


def extract_text(file_bytes: bytes, filename: str) -> ExtractionResult:
    if len(file_bytes) > MAX_SIZE_BYTES:
        return ExtractionResult(
            ok=False,
            error=f"File exceeds the 10MB size cap ({filename}). Please upload a smaller file.",
        )

    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    meta: dict = {}
    try:
        if ext == "pdf":
            raw, meta = _extract_pdf(file_bytes)
        elif ext == "docx":
            raw, meta = _extract_docx(file_bytes)
        else:
            return ExtractionResult(
                ok=False,
                error=f"Unsupported file type '.{ext}'. Please upload a PDF or DOCX file.",
            )
    except Exception as e:
        return ExtractionResult(
            ok=False,
            error=f"Could not parse '{filename}' ({e.__class__.__name__}). "
            "The file may be corrupted or password-protected. Please re-upload.",
        )

    if len(raw.strip()) < MIN_CHARS:
        return ExtractionResult(
            ok=False,
            error=(
                f"Extraction returned very little usable text from '{filename}' "
                "(possibly a scanned or image-based file — this prototype has no "
                "OCR fallback yet). Please re-upload a text-based PDF or DOCX."
            ),
        )

    meta["char_count"] = len(raw.strip())
    return ExtractionResult(ok=True, text=raw, meta=meta)
